import glob
import os
import pandas as pd
from docx import Document
from pathlib2 import Path
from datetime import datetime
import time

# Create list of paths to .doc files
inital_paths = glob.glob('\\Users\\William\\Desktop\\CodeProjects\\MPR2\\*[0-9]', recursive=True)

# Standardizes the paths so they can be sorted by version number, only the most recent MPR is wanted
paths = []
for initial_path in inital_paths:
    presorted_paths = glob.glob(initial_path + '\\**\\[0-9][0-9]*.docx', recursive=True)
    #print(presorted_paths)
    def pathevener(xpath):
        xpath = xpath.replace(" ","")
        xpath = xpath.replace("\Obsolete","")
        xseppath = xpath.split('\\')
        return(xseppath[7])
    sorted_paths = sorted(presorted_paths, key=pathevener, reverse=True)
    try:
        paths.append(sorted_paths[0])
    except:
        continue

output_path = "/Users/William/Desktop/CodeProjects"
writer = pd.ExcelWriter('{}/MPR_tables.xlsx'.format(output_path), engine='xlsxwriter')

excelcounter = 0
yexcelcounter = 0
colcounter = 0
for path in paths:

    time.sleep(1)

    catnum = Path(path).parts[6]
    now = datetime.now()
    current_time = now.strftime("%H:%M:%S")
    print(path, 'Current time =', current_time)
    try:
        # read in word file
        document = Document(path)
        
        #For starting materials table
        
        table = document.tables[3]
        #0 is revised by/Rev. number
        #1 is unknown
        #2 is product technical information
        #3 is for SM
        #4 is for cycle time and yields

        # Data will be a list of rows represented as dictionaries
        # containing each row's data.
        data = []
        
        #keys = ('MPR #', 'Chemical', 'Cat #', 'SKU', 'Costed to process order', 'Den. (liq.)', 'Amt.', 'Mol', 'Mol. Eq.')
        rowcounter = 0
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)
            #print(tuple(text))

            # Establish the mapping based on the first row
            # headers; these will become the keys of our dictionary
            if i == 0:
                continue

            if i == 1:
                keys1 = tuple(text)
                keys2 = ('MPR #',) + keys1 + ('Mol. Eq.',)
                #print(keys2)
                continue
            
            if i > 1:
                text1 = tuple(text)
                text2 = (catnum,) + text1

            # Construct a dictionary for this row, mapping
            # keys to values for this row
            row_data = dict(zip(keys2, text2))
            data.append(row_data)
            rowcounter += 1

        #print(data)
        
        if colcounter == 0:
            df = pd.DataFrame(data)
            df.to_excel(writer, sheet_name='Starting Materials', startrow=excelcounter, index = False)
            excelcounter += rowcounter + 1
        else:
            df = pd.DataFrame(data)
            df.to_excel(writer, sheet_name='Starting Materials', startrow=excelcounter, header = False, index = False)
            excelcounter += rowcounter
        
        #print('excel counter is', excelcounter)
        
        
        #For yield table
        
        ytable = document.tables[4]
        
        x = ytable.rows
        if len(x) > 4:
            ytable = document.tables[5]
        
        ydata = []
        
        ykeys = ('MPR #', 'Direct Hours:', 'Cycle Time:', 'Expected yield:', 'Unit (exp)', 'Theoretical yield:', 'Unit (theo)')
        for j, row in enumerate(ytable.rows):
            ytext = (cell.text for cell in row.cells)
            
            #The expected yield + units
            if j == 0:
                ytextex = tuple(ytext)
                expyld = str(ytextex[0])
                expyldlst = expyld.split()
                truexpyld = expyldlst[5] 
                truexpyldu = expyldlst[6]
                ytext0 = (truexpyld, truexpyldu)
                
                theoyld = str(ytextex[1])
                theoyldlst = theoyld.split()
                trutheoyld = theoyldlst[5]
                trutheoyldu = theoyldlst[6]
                ytext00 = (trutheoyld, trutheoyldu)
                continue
                
            #the expected cycle time
            if j == 1:
                ytextcyc = tuple(ytext)
                expcyc = str(ytextcyc[0])
                expcyclst = expcyc.split()
                truexpcyc = expcyclst[4]
                ytext1 = (truexpcyc,)
                continue
                
            #The direct hours
            if j == 2:
                ytextdirhr = tuple(ytext)
                dirhr = str(ytextdirhr[0])
                dirhrlst = dirhr.split()
                trudirhr = dirhrlst[3]
                ytext2 = (trudirhr,)
                continue
                
            
        ytextfin = (catnum,) + ytext2 + ytext1 + ytext0 + ytext00
            
        # Construct a dictionary for this row, mapping
        # keys to values 
        yrow_data = dict(zip(ykeys, ytextfin))
        ydata.append(yrow_data)

        #print(ydata)
        if colcounter == 0:
            ydf = pd.DataFrame(ydata)
            ydf.to_excel(writer, sheet_name='Yield table', startrow=yexcelcounter, index = False)
            yexcelcounter += 2
        else:
            ydf = pd.DataFrame(ydata)
            ydf.to_excel(writer, sheet_name='Yield table', startrow=yexcelcounter, header = False, index = False)
            yexcelcounter += 1
            
        colcounter += 1
    except Exception as e:
        print('Error in', catnum, e)
        continue
writer.save()

end = datetime.now()
current_time = end.strftime("%H:%M:%S")
print('end time =', current_time)